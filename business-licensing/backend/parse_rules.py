import json
from pathlib import Path
from typing import List, Dict, Any

import docx

RULE_DEFINITIONS = [
    {
        "id": "baseline_registration",
        "title": "דרישות לרישום בסיסי",
        "category": "רישוי",
        "authority": "רשות הרישוי",
        "conditions": {
            "size_m2": {},
            "seating": {},
            "flags": {}
        },
        "keywords": ["רישיון", "רישום"]
    },
    {
        "id": "occupancy_certification",
        "title": "אישור תפוסה",
        "category": "בטיחות",
        "authority": "רשות הכבאות",
        "conditions": {
            "size_m2": {},
            "seating": {"min": 200},
            "flags": {}
        },
        "keywords": ["תפוס", "תפוסה", "אכלוס"]
    },
    {
        "id": "gas_use",
        "title": "שימוש בגז",
        "category": "בטיחות/כיבוי אש",
        "authority": "רשות הכבאות",
        "conditions": {
            "size_m2": {},
            "seating": {},
            "flags": {"uses_gas": True}
        },
        "keywords": ["גז", 'גפ"מ']
    },
    {
        "id": "seating_number",
        "title": "הגבלת מקומות ישיבה",
        "category": "תפעול",
        "authority": "רשות הרישוי",
        "conditions": {
            "size_m2": {},
            "seating": {},
            "flags": {}
        },
        "keywords": ["מקומות ישיבה", "תפוסה"]
    },
    {
        "id": "building_size",
        "title": "גודל מבנה",
        "category": "תכנון ובנייה",
        "authority": "הרשות המקומית",
        "conditions": {
            "size_m2": {"min": 0},
            "seating": {},
            "flags": {}
        },
        "keywords": ["מ\"ר", "שטח"]
    },
    {
        "id": "delivery",
        "title": "משלוח מזון",
        "category": "שירות ללקוח",
        "authority": "משרד הבריאות",
        "conditions": {
            "size_m2": {},
            "seating": {},
            "flags": {"has_delivery": True}
        },
        "keywords": ["משלוח", "שליחת מזון"]
    }
]


def iter_document_elements(doc: docx.Document):
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            yield {
                "kind": "paragraph",
                "index": idx,
                "text": text
            }
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text:
                    yield {
                        "kind": "table",
                        "index": f"{t_idx}:{r_idx}:{c_idx}",
                        "text": text
                    }


def extract_rules(doc_path: Path) -> List[Dict[str, Any]]:
    doc = docx.Document(doc_path)
    elements = list(iter_document_elements(doc))
    results = []
    for rule in RULE_DEFINITIONS:
        matches = [el for el in elements if any(k in el["text"] for k in rule["keywords"])]
        if not matches:
            continue
        first = matches[0]
        result = {
            "id": rule["id"],
            "title": rule["title"],
            "category": rule["category"],
            "authority": rule["authority"],
            "conditions": rule["conditions"],
            "requirements": [m["text"] for m in matches],
            "raw_excerpt": first["text"],
            "provenance": {
                "doc": doc_path.name,
                "section_title": None,
                "heading_level": None,
                "paragraph_index": first["index"] if first["kind"] == "paragraph" else None,
                "table": first["index"] if first["kind"] == "table" else None,
                "page_hint": None
            }
        }
        results.append(result)
    return results


def main():
    import argparse
    parser = argparse.ArgumentParser(description="Extract business licensing rules from docx")
    parser.add_argument("doc", nargs="?", default="data/rules_raw/18-07-2022_4.2A.docx", help="Path to DOCX file")
    parser.add_argument("out", nargs="?", default="data/rules.json", help="Output JSON file")
    args = parser.parse_args()
    path = Path(args.doc)
    rules = extract_rules(path)
    with open(args.out, "w", encoding="utf-8") as f:
        json.dump(rules, f, ensure_ascii=False, indent=2)
    print(f"Extracted {len(rules)} rules to {args.out}")


if __name__ == "__main__":
    main()
